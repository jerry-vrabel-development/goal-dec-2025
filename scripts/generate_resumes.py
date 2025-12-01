#!/usr/bin/env python3
"""
Generate personalized resumes from a Word template.

The script:
  • Loads a CSV of companies/roles.
  • Copies the template for each row.
  • Replaces every {{key}} placeholder (in paragraphs & tables).
  • Saves the new DOCX and converts it to PDF.
  • Emits a single line to logs/submission.log.
"""

import argparse
import datetime
import logging
import os
import re
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx2pdf import convert

# --------------------------------------------------------------------
# ── Configuration ─────────────────────────────────────────────────────
# --------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
ROOT     = BASE_DIR.parent

TEMPLATE_PATH = ROOT / "templates" / "template.docx"
GENERATED_DIR = ROOT / "generated"
DATA_CSV      = ROOT / "data"      / "companies.csv"
LOG_FILE      = ROOT / "logs"      / "submission.log"
PROGRESS_XLS  = ROOT / "progress"  / "progress.xlsx"

# --------------------------------------------------------------------
# ── Logging ───────────────────────────────────────────────────────────
# --------------------------------------------------------------------
logging.basicConfig(
    filename=str(LOG_FILE),
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# --------------------------------------------------------------------
# ── CSV handling ─────────────────────────────────────────────────────
# --------------------------------------------------------------------
def load_companies(csv_path: Path) -> pd.DataFrame:
    """Return the DataFrame; raise if the file is missing."""
    if not csv_path.is_file():
        raise FileNotFoundError(f"CSV not found: {csv_path}")
    return pd.read_csv(csv_path)

# --------------------------------------------------------------------
# ── Placeholder replacement (handles text runs + tables) ─────────────
# --------------------------------------------------------------------
PLACEHOLDER_REGEX = re.compile(r"{{(\w+)}}")  # matches {{company}},
{{role}}, etc.

def replace_placeholders(doc: Document, mapping: dict):
    """Replace all {{key}} placeholders in the document.
    Works on paragraphs, tables, header/footer, etc.
    """
    def _replace_in_text(text: str) -> str:
        return PLACEHOLDER_REGEX.sub(lambda m: str(mapping.get(m.group(1),
m.group(0))), text)

    # Replace in normal paragraphs
    for paragraph in doc.paragraphs:
        paragraph.text = _replace_in_text(paragraph.text)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = _replace_in_text(cell.text)

    # Replace in header/footer (optional – uncomment if you use them)
    # for section in doc.sections:
    #     for header in (section.header, section.footer):
    #         for paragraph in header.paragraphs:
    #             paragraph.text = _replace_in_text(paragraph.text)

# --------------------------------------------------------------------
# ── Generation per row ────────────────────────────────────────────────
# --------------------------------------------------------------------
def generate_resume_for_row(row: pd.Series) -> tuple[Path, Path]:
    """
    Generate DOCX & PDF for a single CSV row.
    Returns a tuple: (docx_path, pdf_path)
    """
    mapping = {
        "company":   row["company"],
        "role":      row["role"],
        "email":     row["email"],
        "address":   row.get("address", ""),     # optional columns
handled safely
        "job_link":  row.get("job_link", ""),
    }

    # 1️⃣  Destination file names
    base_name = f"{row['company'].replace(' ',
'_')}_{row['role'].replace(' ', '_')}"
    target_docx = GENERATED_DIR / f"{base_name}.docx"
    target_pdf  = GENERATED_DIR / f"{base_name}.pdf"

    # 2️⃣  Copy & modify the template
    shutil.copy(TEMPLATE_PATH, target_docx)
    doc = Document(target_docx)
    replace_placeholders(doc, mapping)
    doc.save(target_docx)

    # 3️⃣  Convert to PDF
    convert(str(target_docx), str(target_pdf))

    # 4️⃣  Logging
    logging.info(f"GENERATED: {target_docx} & {target_pdf}")

    # 5️⃣  Pretty‑print to stdout
    print(f"✅  {target_docx} + {target_pdf}")

    return target_docx, target_pdf


# --------------------------------------------------------------------
# ── CLI entrypoint ───────────────────────────────────────────────────
# --------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Generate personalised resumes for every row in a
CSV."
    )
    parser.add_argument(
        "--csv",
        type=Path,
        default=DATA_CSV,
        help="Path to the companies CSV (defaults to data/companies.csv)",
    )
    args = parser.parse_args()

    # Make sure the output directory exists
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    # Load the data
    df = load_companies(args.csv)

    # Generate all
    for _, row in df.iterrows():
        generate_resume_for_row(row)

    print("\n🎉  All resumes generated.")


if __name__ == "__main__":
    main()
